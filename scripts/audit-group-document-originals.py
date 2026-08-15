from __future__ import annotations

import argparse
from pathlib import Path
from zipfile import ZipFile
from xml.etree import ElementTree

from docx import Document


def clean(value: str) -> str:
    return " ".join(value.replace("\xa0", " ").split())


def audit(path: Path) -> None:
    document = Document(path)
    print(f"\n=== {path.name} ===")
    print(f"bytes={path.stat().st_size} sections={len(document.sections)} tables={len(document.tables)}")
    for index, section in enumerate(document.sections, start=1):
        print(
            "section"
            f" {index}: orientation={section.orientation}"
            f" page={section.page_width.inches:.2f}x{section.page_height.inches:.2f}in"
            f" margins={section.left_margin.inches:.2f}/{section.right_margin.inches:.2f}/"
            f"{section.top_margin.inches:.2f}/{section.bottom_margin.inches:.2f}in"
        )
        header = " | ".join(clean(p.text) for p in section.header.paragraphs if clean(p.text))
        footer = " | ".join(clean(p.text) for p in section.footer.paragraphs if clean(p.text))
        if header:
            print(f"header: {header}")
        if footer:
            print(f"footer: {footer}")

    for index, paragraph in enumerate(document.paragraphs, start=1):
        text = clean(paragraph.text)
        if text:
            print(f"P{index:02d} [{paragraph.style.name}]: {text}")

    for table_index, table in enumerate(document.tables, start=1):
        print(f"TABLE {table_index}: rows={len(table.rows)} cols={len(table.columns)}")
        for row_index, row in enumerate(table.rows, start=1):
            values = [clean(cell.text) for cell in row.cells]
            print(f"  R{row_index:02d}: " + " || ".join(values))

    with ZipFile(path) as archive:
        images = [name for name in archive.namelist() if name.startswith("word/media/")]
        headers = [name for name in archive.namelist() if name.startswith("word/header") and name.endswith(".xml")]
        print(f"media={len(images)} headers_xml={len(headers)}")
        for header_name in headers:
            root = ElementTree.fromstring(archive.read(header_name))
            values = [node.text or "" for node in root.iter() if node.tag.endswith("}t")]
            header_text = clean(" ".join(values))
            if header_text:
                print(f"  {header_name}: {header_text}")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("root", type=Path)
    args = parser.parse_args()
    for path in sorted(args.root.rglob("*.docx"), key=lambda item: item.name.casefold()):
        audit(path)


if __name__ == "__main__":
    main()
